#!/usr/bin/env python3
"""Generate GARPIUM-FULL-DOCUMENTATION.docx from project markdown docs."""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUT = DOCS / "GARPIUM-FULL-DOCUMENTATION.docx"

SOURCES = [
    ("Часть I. Продуктовое видение", DOCS / "GARPIUM-PRODUCT-VISION.md"),
    ("Часть II. RBAC и права доступа", DOCS / "GARPIUM-RBAC-ACCESS-MODEL.md"),
    ("Часть III. Implementation Map", DOCS / "GARPIUM-IMPLEMENTATION-MAP.md"),
    ("Часть IV. Итерации разработки", DOCS / "ITERATIONS.md"),
]


def add_md_block(doc: Document, text: str) -> None:
    in_code = False
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("|") and "---" not in line:
            doc.add_paragraph(line.replace("|", " · ").strip(" ·"))
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        else:
            clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            doc.add_paragraph(clean)


def main() -> None:
    doc = Document()
    title = doc.add_heading("GARPIUM LMS", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph("Полная документация платформы — CLOS, RBAC, реализация")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    for part_title, path in SOURCES:
        if not path.exists():
            doc.add_heading(part_title, level=1)
            doc.add_paragraph(f"(Файл {path.name} не найден)")
            doc.add_page_break()
            continue
        doc.add_heading(part_title, level=1)
        add_md_block(doc, path.read_text(encoding="utf-8"))
        doc.add_page_break()

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
